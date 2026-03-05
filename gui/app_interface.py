import os
import shutil
import threading
import time
import customtkinter as ctk
from tkinter import filedialog
from tkinterdnd2 import DND_FILES, TkinterDnD
import fitz  # PyMuPDF
from docx import Document
from PIL import Image

# --- TUS MÓDULOS ---
from core.doc_manager import DocumentManager
from core.ai_assistant import AIAssistantManager
from database.db_handler import DatabaseHandler

# --- CONFIGURACIÓN DE TEMA ---
THEME = {
    "window_bg": ("#EFE7DE", "#0B141A"),
    "sidebar_bg": ("#FFFFFF", "#111B21"),
    "header_bg": ("#F0F2F5", "#202C33"),
    "input_bg": ("#FFFFFF", "#2A3942"),
    "primary": ("#00A884", "#00A884"),
    "primary_hover": ("#06CF9C", "#06CF9C"),
    "primary_disabled": ("#B0E4DD", "#005C4B"),
    "user_bubble": ("#D9FDD3", "#005C4B"),
    "ai_bubble": ("#FFFFFF", "#202C33"),
    "card_bg": ("#E9FCD9", "#025043"),
    "card_top": ("#FFFFFF", "#FFFFFF"),
    "pdf_red": ("#F44336", "#F44336"),
    "text_main": ("#111B21", "#E9EDEF"),
    "text_dim": ("#667781", "#8696A0"),
    "text_white": ("#FFFFFF", "#FFFFFF"),
    "text_card": ("#111B21", "#FFFFFF"),
    "success": ("#00A884", "#00A884"),
    "loading": "#FFD21F"
}

FONTS = {
    "logo": ("Helvetica", 20, "bold"),
    "header": ("Helvetica", 15, "bold"),
    "body": ("Helvetica", 14),
    "meta": ("Helvetica", 11),
    "preview": ("Helvetica", 10, "bold")
}

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("green")

class LAIApp(ctk.CTk, TkinterDnD.DnDWrapper):
    def __init__(self):
        super().__init__()
        self.TkdndVersion = TkinterDnD._require(self)
        self.title("LAI SYSTEM - Auditoría Legal")
        self.geometry("1280x850")
        
        self.db = DatabaseHandler()
        self.doc_manager = DocumentManager()
        self.ai_manager = AIAssistantManager()
        
        self.current_doc_id = None
        self.is_loading = False
        self.typing_active = False 
        self.ver_archivados = False
        
        # --- CARPETA DE ALMACENAMIENTO LOCAL ---
        self.storage_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "storage")
        if not os.path.exists(self.storage_dir):
            os.makedirs(self.storage_dir)
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.setup_ui()
        self.after(200, self.cargar_historial)
        self.mostrar_bienvenida()

    def setup_ui(self):
        self.sidebar = ctk.CTkFrame(self, width=300, corner_radius=0, fg_color=THEME["sidebar_bg"])
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_rowconfigure(3, weight=1)
        
        ctk.CTkLabel(self.sidebar, text="⚖️ LAI SYSTEM", font=FONTS["logo"], text_color=THEME["text_main"]).grid(row=0, column=0, pady=25, padx=20, sticky="w")
        
        self.btn_new = ctk.CTkButton(self.sidebar, text="+ NUEVO CASO", height=45, corner_radius=22, font=FONTS["header"],
                                     fg_color=THEME["primary"], hover_color=THEME["primary_hover"], command=self.reset_chat)
        self.btn_new.grid(row=1, column=0, padx=20, pady=10, sticky="ew")

        self.lbl_list_title = ctk.CTkLabel(self.sidebar, text="CHATS ACTIVOS", font=FONTS["meta"], text_color=THEME["text_dim"])
        self.lbl_list_title.grid(row=2, column=0, sticky="w", padx=25, pady=(20,5))
        
        self.scroll_history = ctk.CTkScrollableFrame(self.sidebar, fg_color="transparent")
        self.scroll_history.grid(row=3, column=0, padx=5, pady=5, sticky="nsew")

        self.footer = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.footer.grid(row=4, column=0, padx=10, pady=20, sticky="ew")
        
        self.btn_archived = ctk.CTkButton(self.footer, text="📂 Ver Archivados", fg_color="transparent", border_width=1, border_color=THEME["text_dim"], text_color=THEME["text_dim"], command=self.alternar_archivados)
        self.btn_archived.pack(side="top", fill="x", pady=5)
        
        self.switch_theme = ctk.CTkSwitch(self.footer, text="Modo Oscuro", command=self.alternar_tema, onvalue="Dark", offvalue="Light")
        self.switch_theme.select()
        self.switch_theme.pack(side="top", pady=10)

        self.main_area = ctk.CTkFrame(self, corner_radius=0, fg_color=THEME["window_bg"])
        self.main_area.grid(row=0, column=1, sticky="nsew")
        self.main_area.grid_rowconfigure(1, weight=1)
        self.main_area.grid_columnconfigure(0, weight=1)

        self.header = ctk.CTkFrame(self.main_area, height=65, fg_color=THEME["header_bg"], corner_radius=0)
        self.header.grid(row=0, column=0, sticky="ew")
        self.lbl_title = ctk.CTkLabel(self.header, text="Mesa de Trabajo", font=FONTS["header"], text_color=THEME["text_main"])
        self.lbl_title.pack(side="left", padx=25)
        self.status_indicator = ctk.CTkLabel(self.header, text="● En línea", font=FONTS["meta"], text_color=THEME["success"])
        self.status_indicator.pack(side="right", padx=25)

        self.chat_feed = ctk.CTkScrollableFrame(self.main_area, fg_color="transparent")
        self.chat_feed.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")

        self.input_frame = ctk.CTkFrame(self.main_area, height=85, fg_color=THEME["header_bg"], corner_radius=0)
        self.input_frame.grid(row=2, column=0, sticky="ew")
        self.input_frame.grid_columnconfigure(2, weight=1)

        self.btn_clip = ctk.CTkButton(self.input_frame, text="+", width=45, height=45, corner_radius=25, 
                                      fg_color=THEME["primary"], hover_color=THEME["primary_hover"], 
                                      text_color="#FFFFFF", font=("Arial", 24), command=self.seleccionar_archivo)
        self.btn_clip.grid(row=0, column=0, padx=(15, 5), pady=15)

        self.btn_gen_doc = ctk.CTkButton(self.input_frame, text="📝", width=45, height=45, corner_radius=25,
                                         fg_color="#3B82F6", hover_color="#2563EB",
                                         text_color="#FFFFFF", font=("Arial", 20), command=self.generar_escrito_legal)
        self.btn_gen_doc.grid(row=0, column=1, padx=(5, 10), pady=15)
        
        self.entry = ctk.CTkEntry(self.input_frame, placeholder_text="Escribe un mensaje...", border_width=0, fg_color=THEME["input_bg"], text_color=THEME["text_main"], font=FONTS["body"], height=48, corner_radius=12)
        self.entry.grid(row=0, column=2, sticky="ew", padx=10, pady=15)
        self.entry.bind("<Return>", lambda e: self.enviar_pregunta())

        self.btn_send = ctk.CTkButton(self.input_frame, text="➤", width=48, height=48, corner_radius=24, fg_color=THEME["primary"], hover_color=THEME["primary_hover"], command=self.enviar_pregunta)
        self.btn_send.grid(row=0, column=3, padx=15)

        self.drop_target_register(DND_FILES)
        self.dnd_bind('<<Drop>>', self.al_soltar_archivo)

    # --- FUNCIONES DE TEMA Y ESTADO ---
    def alternar_tema(self):
        modo = self.switch_theme.get()
        ctk.set_appearance_mode(modo)
        
    def alternar_archivados(self):
        self.ver_archivados = not self.ver_archivados
        estado = "ARCHIVADOS" if self.ver_archivados else "ACTIVOS"
        self.lbl_list_title.configure(text=f"CHATS {estado}")
        self.btn_archived.configure(text="📂 Ver Activos" if self.ver_archivados else "📂 Ver Archivados")
        self.cargar_historial()

    # --- GENERAR ESCRITO ---
    def generar_escrito_legal(self):
        if self.is_loading: return
        dialog = ctk.CTkInputDialog(text="¿Qué documento deseas generar?\n(Ej: Apelación, Contrato)", title="Generar Documento")
        tipo_doc = dialog.get_input()
        if not tipo_doc: return

        self.agregar_mensaje(f"Generar documento: {tipo_doc}", "USUARIO")
        self.set_loading(True, "Redactando documento legal...")
        threading.Thread(target=self._thread_redaccion, args=(tipo_doc,)).start()

    def _thread_redaccion(self, tipo_doc):
        prompt_redaccion = f"ACTUANDO COMO ABOGADO: Basado en la información de este chat y los archivos, redacta un '{tipo_doc}' completo, formal y listo para firmar. Incluye fundamentos de derecho. NO expliques, solo escribe el documento."
        texto_doc = self.ai_manager.enviar_mensaje(prompt_redaccion)
        self.after(0, lambda: self.set_loading(False))
        self.after(0, lambda: self._guardar_word(texto_doc, tipo_doc))

    def _guardar_word(self, texto, tipo):
        try:
            doc = Document()
            doc.add_heading(f'LAI SYSTEM - {tipo.upper()}', 0)
            doc.add_paragraph(texto)
            filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Borrador_{tipo.replace(' ', '_')}.docx")
            if filename:
                doc.save(filename)
                self.agregar_mensaje(f"✅ Documento guardado exitosamente en:\n{filename}", "IA")
            else:
                self.agregar_mensaje("Generación cancelada.", "IA")
        except Exception as e:
            self.agregar_mensaje(f"Error al generar Word: {str(e)}", "IA", "error")

    # --- LÓGICA DE ARCHIVOS (PDF + IMÁGENES) ---
    def obtener_info_archivo(self, path):
        if not path or not os.path.exists(path): return "Archivo no encontrado.", 0, "0 KB", "err"
        
        ext = os.path.splitext(path)[1].lower()
        if ext in ['.jpg', '.jpeg', '.png']:
            peso = os.path.getsize(path)
            peso_str = f"{peso/1024:.0f} KB" if peso < 1024*1024 else f"{peso/(1024*1024):.1f} MB"
            return "IMAGEN", 1, peso_str, "img"
        
        try:
            doc = fitz.open(path)
            texto = doc[0].get_text()[:120].replace('\n', ' ') + "..."
            paginas = len(doc)
            peso = os.path.getsize(path)
            peso_str = f"{peso/1024:.0f} KB" if peso < 1024*1024 else f"{peso/(1024*1024):.1f} MB"
            return texto, paginas, peso_str, "pdf"
        except: return "Archivo protegido.", 0, "0 KB", "err"

    def set_loading(self, active=True, text="Escribiendo..."):
        self.is_loading = active
        if active:
            self.status_indicator.configure(text=f"⏳ {text}", text_color=THEME["loading"])
            self.entry.configure(state="disabled")
            self.btn_send.configure(state="disabled", fg_color=THEME["primary_disabled"])
            self.btn_clip.configure(state="disabled", fg_color=THEME["primary_disabled"])
            self.btn_gen_doc.configure(state="disabled")
        else:
            self.status_indicator.configure(text="● En línea", text_color=THEME["success"])
            self.entry.configure(state="normal")
            self.btn_send.configure(state="normal", fg_color=THEME["primary"])
            self.btn_clip.configure(state="normal", fg_color=THEME["primary"])
            self.btn_gen_doc.configure(state="normal")
            self.entry.focus()

    def agregar_mensaje(self, texto, rol, tipo="texto", animar=False, filepath=None):
        try: self.welcome_frame.destroy()
        except: pass

        row = ctk.CTkFrame(self.chat_feed, fg_color="transparent")
        row.pack(fill="x", pady=8, padx=15)
        align = "right" if rol == "USUARIO" else "left"
        
        if tipo == "archivo":
            info_txt, pags, peso, ftype = self.obtener_info_archivo(filepath)
            
            card = ctk.CTkFrame(row, fg_color=THEME["card_bg"], corner_radius=12, width=320)
            card.pack(side=align)

            top = ctk.CTkFrame(card, fg_color=THEME["card_top"], corner_radius=10, height=100)
            top.pack(fill="x", padx=4, pady=(4, 0))
            top.pack_propagate(False)
            
            if ftype == "img":
                try:
                    img = ctk.CTkImage(Image.open(filepath), size=(280, 80))
                    ctk.CTkLabel(top, text="", image=img).pack(expand=True, fill="both")
                except:
                    ctk.CTkLabel(top, text="📸 IMAGEN", font=FONTS["preview"], text_color="#667781").pack(anchor="center", pady=30)
            else:
                ctk.CTkLabel(top, text="VISTA PREVIA", font=FONTS["preview"], text_color="#667781").pack(anchor="w", padx=12, pady=(10, 2))
                ctk.CTkLabel(top, text=info_txt, font=("Helvetica", 11), text_color="#333", wraplength=280, justify="left").pack(anchor="w", padx=12)

            bot = ctk.CTkFrame(card, fg_color="transparent")
            bot.pack(fill="x", padx=10, pady=10)
            
            color_icon = "#3B82F6" if ftype == "img" else THEME["pdf_red"]
            txt_icon = "IMG" if ftype == "img" else "PDF"
            
            icon = ctk.CTkFrame(bot, fg_color=color_icon, width=38, height=45, corner_radius=5)
            icon.pack(side="left")
            icon.pack_propagate(False)
            ctk.CTkLabel(icon, text=txt_icon, font=("Arial", 10, "bold"), text_color="white").place(relx=0.5, rely=0.5, anchor="center")
            
            info = ctk.CTkFrame(bot, fg_color="transparent")
            info.pack(side="left", padx=12)
            
            ctk.CTkLabel(info, text=texto[:25], font=("Helvetica", 13, "bold"), text_color=THEME["text_white"]).pack(anchor="w")
            sub_info = f"{peso} • {ftype.upper()}" if ftype == "img" else f"{pags} pág • {peso} • PDF"
            ctk.CTkLabel(info, text=sub_info, font=FONTS["meta"], text_color=THEME["text_dim"]).pack(anchor="w")
        else:
            bubble = ctk.CTkFrame(row, fg_color=THEME["user_bubble"] if rol=="USUARIO" else THEME["ai_bubble"], corner_radius=15)
            bubble.pack(side=align, padx=5)
            lbl = ctk.CTkLabel(bubble, text=texto if not animar else "", font=FONTS["body"], text_color=THEME["text_main"], wraplength=650, justify="left")
            lbl.pack(padx=15, pady=10)
            if animar:
                self.typing_active = True
                self._typewriter(lbl, texto)
        
        self.chat_feed._parent_canvas.yview_moveto(1.0)

    def _typewriter(self, label, full_text, idx=0):
        try:
            if not self.typing_active or idx >= len(full_text):
                self.typing_active = False
                label.configure(text=full_text)
                return
            label.configure(text=full_text[:idx+1])
            self.after(5, lambda: self._typewriter(label, full_text, idx+1))
        except Exception:
            return

    # --- INTERACCIÓN ---
    def enviar_pregunta(self):
        txt = self.entry.get()
        if not txt.strip() or self.is_loading: return
        self.entry.delete(0, "end")
        self.agregar_mensaje(txt, "USUARIO")
        
        msg_load = "Consultando Constitución..." if "ley" in txt.lower() else "Analizando respuesta..."
        self.set_loading(True, msg_load)
        
        if not self.current_doc_id:
            tid = self.ai_manager.nuevo_hilo()
            self.current_doc_id = self.db.registrar_documento("Consulta", tid)
            self.after(0, self.cargar_historial)

        self.db.guardar_mensaje(self.current_doc_id, "user", txt)
        threading.Thread(target=self._hilo_chat, args=(txt,)).start()

    def _hilo_chat(self, txt):
        resp = self.ai_manager.enviar_mensaje(txt)
        self.after(0, lambda: self.set_loading(False))
        self.after(0, lambda: self.agregar_mensaje(resp, "IA", animar=True))
        self.db.guardar_mensaje(self.current_doc_id, "assistant", resp)

    def al_soltar_archivo(self, event):
        if self.is_loading: return
        path = event.data.replace("{", "").replace("}", "")
        self.procesar_archivo(path)

    def seleccionar_archivo(self):
        if self.is_loading: return
        path = filedialog.askopenfilename(filetypes=[("Documentos e Imágenes", "*.pdf *.jpg *.jpeg *.png")])
        if path: self.procesar_archivo(path)

    def procesar_archivo(self, path):
        # 1. COPIA LOCAL PARA PERSISTENCIA
        try:
            filename = os.path.basename(path)
            unique_name = f"{int(time.time())}_{filename}"
            dest_path = os.path.join(self.storage_dir, unique_name)
            shutil.copy2(path, dest_path) 
            local_path = dest_path 
        except Exception as e:
            print(f"Error copiando archivo: {e}")
            local_path = path 

        self.agregar_mensaje(filename, "USUARIO", "archivo", filepath=local_path)
        self.set_loading(True, "Analizando archivo...")
        self.lbl_title.configure(text=f"Caso: {filename[:20]}...")
        
        ext = os.path.splitext(local_path)[1].lower()
        if ext in ['.jpg', '.jpeg', '.png']:
            threading.Thread(target=self._thread_upload_imagen, args=(local_path,)).start()
        else:
            threading.Thread(target=self._thread_upload, args=(local_path,)).start()

    def _thread_upload(self, path):
        if not self.current_doc_id:
            tid = self.ai_manager.nuevo_hilo()
            self.current_doc_id = self.db.registrar_documento(os.path.basename(path), tid, "Análisis")
            self.after(0, self.cargar_historial)

        exito, resp = self.ai_manager.procesar_archivo(path)
        self.after(0, lambda: self.set_loading(False))
        if exito:
            self.after(0, lambda: self.agregar_mensaje(resp, "IA", animar=True))
            self.db.guardar_mensaje(self.current_doc_id, "user", f"[ARCHIVO] {path}")
            self.db.guardar_mensaje(self.current_doc_id, "assistant", resp)
        else:
            self.after(0, lambda: self.agregar_mensaje(resp, "IA", "error"))

    def _thread_upload_imagen(self, path):
        if not self.current_doc_id:
            tid = self.ai_manager.nuevo_hilo()
            self.current_doc_id = self.db.registrar_documento(os.path.basename(path), tid, "Análisis Visual")
            self.after(0, self.cargar_historial)

        resp = self.ai_manager.enviar_mensaje("Analiza esta imagen legal/evidencia que acabo de subir:", imagen_path=path)
        
        self.after(0, lambda: self.set_loading(False))
        self.after(0, lambda: self.agregar_mensaje(resp, "IA", animar=True))
        
        self.db.guardar_mensaje(self.current_doc_id, "user", f"[IMAGEN] {path}")
        self.db.guardar_mensaje(self.current_doc_id, "assistant", resp)

    def cargar_historial(self):
        for w in self.scroll_history.winfo_children(): w.destroy()
        if self.ver_archivados:
            docs = self.db.obtener_archivados()
        else:
            docs = self.db.obtener_activos()
        if not docs: return

        for doc in docs:
            f = ctk.CTkFrame(self.scroll_history, fg_color="transparent")
            f.pack(fill="x", pady=2)
            titulo = doc.get('resumen_ia') or doc.get('nombre_archivo')
            
            btn = ctk.CTkButton(f, text=f"📄 {titulo[:18]}...", fg_color="transparent", anchor="w", text_color=THEME["text_dim"], hover_color=THEME["input_bg"], command=lambda d=doc: self.restaurar_sesion(d))
            btn.pack(side="left", fill="x", expand=True)
            
            btn_edit = ctk.CTkButton(f, text="✎", width=25, fg_color="transparent", text_color="#AAA", hover_color="#333", command=lambda d=doc: self.renombrar_chat(d))
            btn_edit.pack(side="right")
            
            ico_arch = "📥" if self.ver_archivados else "📂"
            cmd_arch = lambda d=doc: [self.db.archivar_documento(d['id'], not self.ver_archivados), self.cargar_historial()]
            btn_arch = ctk.CTkButton(f, text=ico_arch, width=25, fg_color="transparent", text_color="#AAA", hover_color="#333", command=cmd_arch)
            btn_arch.pack(side="right")
            
            btn_del = ctk.CTkButton(f, text="🗑", width=25, fg_color="transparent", text_color="#E53935", hover_color="#330000", command=lambda d=doc: [self.db.eliminar_definitivamente(d['id']), self.cargar_historial()])
            btn_del.pack(side="right")

    def renombrar_chat(self, doc):
        dialog = ctk.CTkInputDialog(text="Nuevo nombre del caso:", title="Renombrar")
        nuevo = dialog.get_input()
        if nuevo:
            self.db.actualizar_titulo(doc['id'], nuevo) 
            self.cargar_historial()

    def restaurar_sesion(self, doc):
        self.current_doc_id = doc['id']
        self.ai_manager.set_thread(doc['thread_id'])
        self.lbl_title.configure(text=f"Caso: {doc.get('resumen_ia')[:25]}")
        self.limpiar_chat()
        try: self.welcome_frame.destroy()
        except: pass
        
        for m in self.db.recuperar_chat_completo(doc['id']):
            rol = "USUARIO" if m['rol'] == "user" else "IA"
            contenido = m['contenido']
            
            if "[ARCHIVO]" in contenido:
                path = contenido.replace("[ARCHIVO] ", "")
                nombre_mostrar = os.path.basename(path)
                self.agregar_mensaje(nombre_mostrar, rol, "archivo", filepath=path)
            
            elif "[IMAGEN]" in contenido:
                path = contenido.replace("[IMAGEN] ", "")
                nombre_mostrar = os.path.basename(path)
                self.agregar_mensaje(nombre_mostrar, rol, "archivo", filepath=path)
            
            else:
                self.agregar_mensaje(contenido, rol, "texto")

    def mostrar_bienvenida(self):
        self.limpiar_chat()
        self.welcome_frame = ctk.CTkFrame(self.chat_feed, fg_color="transparent")
        self.welcome_frame.pack(expand=True, fill="both", pady=80)
        ctk.CTkLabel(self.welcome_frame, text="⚖️", font=("Arial", 80)).pack(pady=20)
        ctk.CTkLabel(self.welcome_frame, text="LAI System", font=("Helvetica", 32, "bold"), text_color=THEME["text_main"]).pack()
        ctk.CTkLabel(self.welcome_frame, text="Tu asistente legal inteligente", font=FONTS["body"], text_color=THEME["text_dim"]).pack(pady=10)

    def limpiar_chat(self):
        for w in self.chat_feed.winfo_children(): w.destroy()

    def reset_chat(self):
        self.current_doc_id = None
        self.ai_manager.nuevo_hilo()
        self.limpiar_chat()
        self.lbl_title.configure(text="Nueva Auditoría")
        self.mostrar_bienvenida()

if __name__ == "__main__":
    app = LAIApp()
    app.mainloop()